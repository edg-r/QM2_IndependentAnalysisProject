from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


ROOT = Path("/Users/edgar/Documents/01 Projects/GPCO 454 - QM2 - Ravanilla/IAP")
SOURCE = ROOT / "IAP Write Up.docx"
OUTPUT = ROOT / "IAP Write Up - Cross Section Revisions Highlighted Minimal.docx"


def clear_paragraph(paragraph):
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def set_highlighted_paragraph_text(paragraph, text):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_highlighted_cell_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        if not paragraph.runs:
            run = paragraph.add_run("")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            continue
        for run in paragraph.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def revise_paragraphs(doc):
    replacements = {
        39: (
            "\tWith a large US pull back on foreign aid contributions in the UN and otherwise, many "
            "in the international community look to China to fill this gap. Many see it logical for "
            "China to fill this gap in influence, however traditionally China has been selective in "
            "where it allocates its ODA. This research was designed to answer if China’s foreign aid "
            "is tightly correlated to more authoritarian governments. China’s aid allocation differs "
            "from the US which is typically concentrated in the UN, while China’s is more direct and "
            "thus bilateral. Chinas development finance and belt and road initiative is seen more "
            "politically motivated in the West, when compared to Western aid. Therefor the purpose of "
            "this research is to see if China’s Aid disbursement method has preference to other "
            "authoritarian governments. After extensive quantitative research using a country level "
            "cross sectional regression analysis, there is a positive association. However, after "
            "economic and demographic controls are accounted for in the full sample, the countries "
            "government does not have a statistically significant relationship with how much aid a "
            "country receives. At the same time, among countries that did receive positive Chinese aid, "
            "a one point increase in autocracy score is associated with about 2.95 times higher mean "
            "annual aid. Suggesting China may have different criterion for aid disbursement once aid "
            "is already being given."
        ),
        41: "Our research question is, do more autocratic regimes receive more funding from China on average?",
        42: (
            "Null Hypothesis: Autocratic Regimes do not differ in average aid received when compared "
            "to Democratic ones."
        ),
        43: (
            "Alternative Hypothesis: Autocratic regimes receive more aid on average when compared to "
            "Democratic ones."
        ),
        46: (
            "\tUsing the Our World in Data (OWID) data set for classification of autocratic governments "
            "and the Chinese Global Development Finance Data set for levels of foreign Aid provided by "
            "China to run my cross sectional regression. Our final data set collapsed the years of 2013 "
            "to 2021 into country averages, making the unit of analysis the country. The bivariate model "
            "utilized 174 countries, the controls model utilized 135 countries, and the positive aid only "
            "model utilized 92 countries. Our reconfigured OWID data set classified autocratic governments "
            "on a scale from 0 to 3, with 0 being electoral democracies and 3 being closed autocracies. "
            "China itself is classified as a closed autocracy. In the graph below we can see that even "
            "when adjusting aid as a logarithm of its mean annual value so that extremely large dollar "
            "donations do not dominate the graph, more autocratic regimes received more aid from China. "
            "We also see that after a regression analysis there is a positive relationship between "
            "autocracy levels and level of aid provided in the bivariate model, though in the full "
            "controls model this relationship loses statistical significance."
        ),
        50: (
            "Our regression structure is a simple cross sectional linear regression. After controlling "
            "for variables this positive relationship loses statistical significance in the full sample. "
            "For example, GDP per capita to control for development levels of economies and their "
            "correlated effects on governmental structures. Corruption index was used to control for "
            "government quality, as perhaps China utilizes it to assess efficacy of their aid which may "
            "be lost to corruption. Extreme poverty share acts as a control for aid need in a country. "
            "As countries with higher percentage of the population living in extreme poverty are in "
            "greater need of aid. Gini coefficient controlled for level of inequality within a country "
            "receiving aid. Finally, because the data was collapsed into country averages from 2013 to "
            "2021, the model compares countries on their mean aid levels rather than controlling for "
            "country specific and year specific changes in aid disbursement. All this to isolate any "
            "correlational relationship between style of government and the corresponding level of "
            "Chinese aid. After all controls were taken into account in the full sample, it is clear "
            "that there is no statistically relevant correlational relationship between level of "
            "autocracy and amount of Chinese aid. However, among countries that did receive aid, the "
            "autocracy score remains statistically significant, and a one point increase in autocracy "
            "score is associated with about 2.95 times higher mean annual aid. These findings show that "
            "within the full sample China does not appear to allocate aid differently by regime type, "
            "but among recipient countries autocracy still appears to matter."
        ),
        52: (
            "\tSome limitations of research were that GDP of countries could not be taken into account "
            "as only GDP per Capita was provided within the data set. This could introduce a bias for "
            "larger populations or economies which China may be more inclined to provide aid or shy away "
            "from providing aid. Another limitation is that by collapsing the data into a country level "
            "cross section, year to year variation is no longer captured. The full controls model also "
            "loses observations because of missing control variables, and some country cases appear to "
            "have stronger influence on the estimates than others. Therefore, while the full sample does "
            "not show a statistically significant relationship between autocracy and Chinese aid, the "
            "positive aid only findings should be interpreted more cautiously."
        ),
    }

    for index, text in replacements.items():
        set_highlighted_paragraph_text(doc.paragraphs[index], text)


def revise_table(doc):
    table = doc.tables[1]

    updates = {
        (0, 0): "Chinese Aid and Regime Type (Country-Level Cross-Section)",
        (0, 1): "Chinese Aid and Regime Type (Country-Level Cross-Section)",
        (0, 2): "Chinese Aid and Regime Type (Country-Level Cross-Section)",
        (0, 3): "Chinese Aid and Regime Type (Country-Level Cross-Section)",
        (4, 1): "Log(1 + mean annual Chinese aid in constant USD 2021)",
        (4, 2): "Log(1 + mean annual Chinese aid in constant USD 2021)",
        (4, 3): "Log(1 + mean annual Chinese aid in constant USD 2021)",
        (5, 1): "Bivariate",
        (5, 2): "Controls",
        (5, 3): "Controls, positive-aid only",
        (8, 0): "Autocracy score",
        (8, 1): "4.186***",
        (8, 2): "0.948",
        (8, 3): "1.083***",
        (9, 1): "(0.769)",
        (9, 2): "(1.171)",
        (9, 3): "(0.397)",
        (11, 0): "Log GDP per capita",
        (11, 1): "",
        (11, 2): "-2.591***",
        (11, 3): "0.812***",
        (12, 1): "",
        (12, 2): "(0.960)",
        (12, 3): "(0.276)",
        (14, 0): "Corruption Perceptions Index",
        (14, 1): "",
        (14, 2): "-0.207***",
        (14, 3): "-0.009",
        (15, 1): "",
        (15, 2): "(0.061)",
        (15, 3): "(0.026)",
        (17, 0): "Extreme poverty share",
        (17, 1): "",
        (17, 2): "-0.111**",
        (17, 3): "0.022",
        (18, 1): "",
        (18, 2): "(0.046)",
        (18, 3): "(0.013)",
        (20, 0): "Gini index",
        (20, 1): "",
        (20, 2): "29.888***",
        (20, 3): "1.008",
        (21, 1): "",
        (21, 2): "(8.620)",
        (21, 3): "(3.314)",
        (23, 0): "Constant",
        (23, 1): "7.642***",
        (23, 2): "35.808***",
        (23, 3): "9.717***",
        (24, 1): "(1.165)",
        (24, 2): "(9.093)",
        (24, 3): "(2.533)",
        (25, 0): "Unit of analysis",
        (25, 1): "Country",
        (25, 2): "Country",
        (25, 3): "Country",
        (27, 0): "Years collapsed",
        (27, 1): "2013-2021 mean",
        (27, 2): "2013-2021 mean",
        (27, 3): "2013-2021 mean",
        (28, 0): "Sample",
        (28, 1): "All countries",
        (28, 2): "All countries",
        (28, 3): "Aid recipients only",
        (29, 0): "Standard errors",
        (29, 1): "HC1",
        (29, 2): "HC1",
        (29, 3): "HC1",
        (30, 0): "Observations",
        (30, 1): "174",
        (30, 2): "135",
        (30, 3): "92",
        (31, 0): "R2",
        (31, 1): "0.202",
        (31, 2): "0.594",
        (31, 3): "0.165",
        (32, 0): "Adjusted R2",
        (32, 1): "0.197",
        (32, 2): "0.578",
        (32, 3): "0.117",
    }

    for (row_index, col_index), value in updates.items():
        set_highlighted_cell_text(table.cell(row_index, col_index), value)


def main():
    doc = Document(SOURCE)
    revise_paragraphs(doc)
    revise_table(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
